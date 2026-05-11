import docx
import pandas as pd
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from bs4 import BeautifulSoup  # Import BeautifulSoup for parsing HTML
import time
import json
import openai
from fake_useragent import UserAgent
import undetected_chromedriver as uc  # For better detection avoidance
import numpy as np
import random
from selenium.webdriver.chrome.options import Options


def initialize_driver():
    """Initialize the Selenium WebDriver with required options."""
    chrome_options = Options()
    chrome_options.add_argument("--headless")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1920,1080")
    driver = webdriver.Chrome(options=chrome_options)
    return driver

# Load the Word document
# doc = docx.Document(r"C:\Users\jagadesw.devarasetti\Downloads\web_scrapping (19).docx") #
def web_scrapping_all(client,doc,tableA_df):

    with open(r'Config\configuration.json', 'r') as f:
        config = json.load(f)
    # Set the OpenAI API key
    openai.api_key = config["api_key"]
    all_tokens = []

    # Extract the tables from the document
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text)
            table_data.append(row_data)
        tables.append(table_data)

    # Convert the extracted tables into a DataFrame
    # df = pd.DataFrame(tables[1]) 
    # Determine which table to process based on tableA_df
    df = pd.DataFrame(tables[1]) if tableA_df is not None else pd.DataFrame(tables[0])
    update_table_index = 1 if tableA_df is not None else 0

    # Drop rows where any cell contains "data not found"
    df_cleaned = df[~df.apply(lambda row: row.astype(str).str.contains("data not found", case=False).any(), axis=1)]

    # Update the table in the Word document
    def update_table2_in_doc(doc, df):
        # Assuming the table to update is the second table (index 1)
        table = doc.tables[update_table_index]

        # Remove extra rows from the Word table
        while len(table.rows) > len(df):
            table._element.remove(table.rows[-1]._element)

        # Add rows if the DataFrame has more rows than the table
        while len(table.rows) < len(df):
            table.add_row()

        # Update the table content with the cleaned DataFrame
        for i, row in enumerate(df.values):
            for j, value in enumerate(row):
                table.cell(i, j).text = str(value)

    # Update the table in the document with the cleaned DataFrame
    update_table2_in_doc(doc, df_cleaned)


    # Convert the extracted tables into a DataFrame
    df = pd.DataFrame(tables[2]) if tableA_df is not None else pd.DataFrame(tables[1]) 

    df.columns = df.iloc[0]
    df = df.iloc[0:]
    df = df.drop(df.index[0])

    # Creating the desired output dictionary
    df.set_index("Parameters", inplace=True)
    # Update the 0th row's cell values with column names, excluding the 0th, 1st, and last columns
    for col_index in range(1, len(df.columns) - 1):  # Start from index 2 and exclude the last column
        df.iloc[0, col_index] = df.columns[col_index]  # Update the 0th row's cell with the column name

    # Define a function to extract text after "Class B"
    def extract_after_class_b(text):
        if isinstance(text, str) and "Class B" in text:
            match = re.search(r'Class B\s*(.*)', text)
            if match:
                extracted_text = match.group(1).strip()
                # Check if extracted text is empty, NaN, or other null values
                if extracted_text in ["", "NaN", "nan", "NULL", "null", None, np.nan]:
                    return text  # Keep original value
                return extracted_text  # Return extracted value
        return text  # If "Class B" not found, return original text

    # Exclude first two and last columns
    cols_to_update = df.columns[1:-1]  
    # print(cols_to_update)

    # Apply transformation to the 0th row
    df.iloc[0, 1:-1] = df.iloc[0, 1:-1].apply(extract_after_class_b)

    # print(df)
    # print("****************************************")
    # print("****************************************")


    output = {} 

    # for param in df.index:
    #     for col in df.columns:
    #         if df.loc[param, col] == "Data not found":
    #             output.setdefault(param, []).append(col)

    for param in df.index:
        for col in df.columns:
            # print(df.loc[param, col])
            # print(type(df.loc[param, col]))
            print(f"Checking param: {param}, col: {col}")
            # Ensure single-value access
            cell_value = df.loc[param, col]  # Access the cell value
            if isinstance(cell_value, pd.Series):  # If multiple values are returned
                for value in cell_value:
                    if value == "Data not found":
                        output.setdefault(param, []).append(col)
            else:  # Single value
                if cell_value == "Data not found":
                    output.setdefault(param, []).append(col)

    def search_google(query):
        # Set up Chrome WebDriver
        ua = UserAgent()  # Initialize UserAgent
        chrome_options = uc.ChromeOptions()
        chrome_options.add_argument(f"user-agent={ua.random}")  # Random user agent
        # chrome_options.add_argument("--headless")
        driver = uc.Chrome(driver_executable_path=ChromeDriverManager().install(), options=chrome_options)
        # driver = initialize_driver()
        # Open Google
        driver.get("https://www.google.com")

        # Find the search bar and enter the query
        search_box = driver.find_element(By.NAME, "q")
        search_box.send_keys(query)
        search_box.send_keys(Keys.RETURN)

        time.sleep(5)  # Wait for results to load

        # Extract all text from the page
        page_source = driver.page_source  # Get the full HTML of the page
        soup = BeautifulSoup(page_source, 'html.parser')  # Parse the HTML using BeautifulSoup
        all_text = soup.get_text(separator=' ')  # Extract all text from the page
        # print(all_text)
        driver.quit()
        return all_text  # Return the extracted text


    def extract_attribute(client,query, attribute):
        """
        Generic function to extract specific attributes from the web data.
        """
        web_data = search_google(query)

        # Define prompts for different attributes
        prompts = {
            "Storage Volume": "Please extract only storage volume for the device. Don't return entire text or statement. Just return the storage volume. Example: 100L or 300 to 400 liters in this format.",
            "Temperature Range": "Please extract only Temperature Range for the device. Don't return entire text or statement. Just return the Temperature Range. Example: -50°C to -80°C or -60°C.",
            "Orientation": "Please extract only Orientation for the device. Don't return entire text or statement. Just return the Orientation. Example: Upright, vertical, horizontal, etc.",
            "Voltage": "Please extract only Voltage for the device. Don't return entire text or statement. Just return the Voltage. Example: 100V or 200V or 230 / 240V, 50 Hz 230V, 50 Hz, 1 Phase.",
            "Material": "Please extract only Material for the device. Don't return entire text or statement. Just return the Material. Example: Interior: Stainless Steel / Exterior: Painted Steel.",
            "Insulation": "Please extract only Insulation for the device. Don't return entire text or statement. Just return the Insulation. Example: High-density CFC-free polyurethane foam insulation or High-density polyurethane foam.",
            "Refrigeration system": "Please extract only Refrigeration system for the device. Don't return entire text or statement. Just return the Refrigeration system. Example: Cascade refrigeration system or Independent Dual-Cooling or Independent Cascade Cooling",
            "Alarms": "Please extract only Alarms for the device. Don't return entire text or statement. Just return the Alarms. Example: Audible and visual alarms for temperature fluctuations and power failure, Adjustable Warm and cold alarms, Power Failure Alarms, Clean filter Alarm, Power failure alarm, Battery Low alarm, Door open alarm etc.",    
            "Indoor/Outdoor Usage": "Please extract only Indoor/Outdoor Usage for the device. Don't return entire text or statement. Just return the Indoor/Outdoor Usage. Example: Indoor or Outdoor or Indoor/Outdoor"
        
        
        }

        # Get the appropriate prompt for the attribute
        prompt = f"""
        {query}

        Here is some additional information from Google search:
        {web_data}

        {prompts.get(attribute, "Please extract the required information.")}
        """

        # Call OpenAI API
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": f"You are an expert in extracting {attribute.lower()} from text."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.0
        )
        response = response.model_dump()
        return response['choices'][0]['message']['content']

    data = output
    # Remove 'License Holder' key
    data.pop('License Holder', None)

    # Extract unique devices
    devices = set(device for values in data.values() for device in values)

    # Generate queries
    queries = []
    for device in devices:
        for attribute, devices_list in data.items():
            if device in devices_list:
                # if attribute == "Storage Volume" or attribute == "Temperature Range":

                query=f"What is the {attribute} for '{device}' device?"
                res = extract_attribute(client,query, attribute)  # Use the generic function
                print(f"Result for {device} - {attribute}: {res}")

                # print(res)
                if attribute in df.index and device in df.columns:
                    df.loc[attribute, device] = res  # Update the specific cell in the DataFrame

    # Print the updated DataFrame
    # print(df)
    # Function to update the table in the Word document with the DataFrame
    def update_table3_in_doc(doc, df,update_table3_index):
        # Assuming the table to update is the third table (index 2)
        table = doc.tables[update_table3_index]

        # Clear the existing table content
        for row in table.rows:
            for cell in row.cells:
                cell.text = ""

        # Add the updated DataFrame content to the table
        # Add the header row
        for col_idx, col_name in enumerate(df.columns, start=1):
            table.cell(0, col_idx).text = col_name

        # Add the rows from the DataFrame
        for row_idx, (index, row) in enumerate(df.iterrows(), start=1):
            table.cell(row_idx, 0).text = str(index)  # Add the index (Parameters)
            for col_idx, value in enumerate(row, start=1):
                table.cell(row_idx, col_idx).text = str(value)
        # print(df)
        print("****************************************")
        # print(table)
    def Equivalence_Generation(client,row):
        # prompt = f"""
        #  Analyze the following search row: {row} about a parameter, The first value contains the name of a parameter. The other values contain
        #  the information about that parameter for different device series, compare them among each other and give a one line comparison of at max 20 words. 
        #  Mention the similarities and differences just about the parameter dont include brand names or external information in your comparision. 
        #  Dont compare the Parameters that dont make sense to be compared in a device comparison table like License Holder, License Number, etc.
        #  Comparison should be apple to apple comparison, dont compare different traits between them. 
        #  There should be strictly no model response. 
        # """
        prompt = f"""
        Analyze the following row: {row}. The first value contains the name of a parameter. The other values contain
        information about that parameter for different device series. Compare them among each other and provide a one-line
        summary of at most 20 words. If the context of all values is the same (even if phrased differently), return "Similar".
        If there are differences in the context or meaning, describe the differences concisely without including brand names
        or external information. Focus only on the parameter's similarities and differences. Do not compare unrelated traits
        or parameters that don't make sense to compare, such as License Holder or License Number. Ensure the comparison is
        clear and concise.
        """

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert in competitor product analysis."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.2

        )
        start_time = time.time()
        response = response.model_dump()
        # Return the generated content if available
        if response and response['choices']:
            input_tokens = response['usage']['prompt_tokens']
            output_tokens = response['usage']['completion_tokens']
            total_tokens = response['usage']['total_tokens']
            all_tokens.append({
                'input_tokens': input_tokens,
                'output_tokens': output_tokens,
                'total_tokens': total_tokens,
                'response_time': time.time() - start_time
            })
        return response['choices'][0]['message']['content']

    
    update_table3_index = 2 if tableA_df is not None else 1
    # Update the table in the document with the latest DataFrame
    update_table3_in_doc(doc, df,update_table3_index)
    final_df = df
    equivalence=[]
    for row in final_df.values:
        # print(row)
        equivalence_string = Equivalence_Generation(client,row)
        equivalence.append(equivalence_string)

    final_df['Equivalence']=equivalence
    tableC_df = final_df
    # print(tableC_df)

    # tableC_df.to_excel(r"C:\Users\jagadesw.devarasetti\Downloads\tableC.xlsx", index=False)
    # Save the updated document
    doc.save(r"updated_web_scrapping2104.docx")
    return doc

# from docx import Document
# # Call the function to perform web scraping and update the document
# doc_path = r"C:\Users\jagadesw.devarasetti\OneDrive - Thermo Fisher Scientific\Desktop\DocumentGenerator-development\web_scrapping_tables_final.docx"
# doc = Document(doc_path)
# tableA_df ="hgsgjs" #None
# web_scrapping_all(doc,tableA_df)