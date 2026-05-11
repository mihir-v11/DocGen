import io
from docx import Document
import base64
import docx
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import pandas as pd
from docx.shared import Inches, RGBColor,Pt
from docx.oxml.ns import nsdecls
import io
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_ALIGN_PARAGRAPH
import streamlit as st














def set_table_border(table):
    # Loop through rows and cells to set borders for each cell
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            borders = OxmlElement('w:tcBorders')  # Create table cell borders element
            
            # Define border styles (top, bottom, left, right)
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')  # Border type: single line
                border.set(qn('w:sz'), '4')       # Border size: 4 (1/8 pt)
                border.set(qn('w:space'), '0')    # No space between border and content
                border.set(qn('w:color'), '000000')  # Border color: black
                borders.append(border)  # Append border element
            
            tcPr.append(borders)  # Add borders to cell properties

def set_cell_background(cell, color):
    """Set cell background shading color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Create a new shading element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)  # Set fill color
    tcPr.append(shd)


def _handle_bullet_line(doc, line, bullet_symbols, in_break=False, cap_symbol=False):
    """
    Create a bulleted paragraph from `line`.
    - Accepts real tabs '\t' or literal '\\t'.
    - Counts leading tabs -> tab_count.
    - Removes all leading '*' characters.
    - Chooses symbol: if cap_symbol True use last symbol for deeper indents,
      otherwise cycle using modulo.
    - Preserves inline bold **...**.
    """
    # Normalize literal "\t" to actual tab
    normalized = line.replace("\\t", "\t")

    # Count leading real tabs
    tab_count = 0
    while tab_count < len(normalized) and normalized[tab_count] == "\t":
        tab_count += 1

    # Remove leading tabs and then any leading spaces
    without_tabs = normalized[tab_count:].lstrip(" ")

    # Skip ALL leading '*' characters (handles '*', '**', '***', or malformed '* *')
    idx = 0
    while idx < len(without_tabs) and without_tabs[idx] == "*":
        idx += 1
    clean_line = without_tabs[idx:].lstrip()

    # Pick symbol
    if cap_symbol:
        symbol = bullet_symbols[min(tab_count, len(bullet_symbols) - 1)]
    else:
        # new (1-tab -> first symbol, caps at last for deeper)
        index = max(0, tab_count - 1)
        symbol = bullet_symbols[min(index, len(bullet_symbols) - 1)]


    p = doc.add_paragraph()
    # indent amount per level (adjust multiplier as needed)
    p.paragraph_format.left_indent = Inches(0.3 * tab_count)

    # handle inline bold segments using the same regex you used before
    if "**" in clean_line:
        parts = re.split(r"(\*\*.*?\*\*)", clean_line)
        p.add_run(f"{symbol} ")
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
    else:
        p.add_run(f"{symbol} {clean_line}")

    # apply the same "in_break" font sizing behavior you used
    if in_break:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(24)


def process_text_to_docx(doc, text, bullet_symbols=None, cap_bullet_symbol=False):
    """
    Unified processor for your input text -> DOCX.
    - Supports: <break>, ## headings, **bold** headings/inline, bullets (* & \t*),
      multi-tab indents using \t or \\t, and default paragraphs.
    - bullet_symbols: list of symbols (default ["•","○","▪"])
    - cap_bullet_symbol: if True, deeper indents use last symbol (no cycling)
    """
    if bullet_symbols is None:
        bullet_symbols = ["•", "○", "▪"]

    lines = text.splitlines()
    in_break = False
    i = 0

    def set_paragraph_font(paragraph):
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(24)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        # debug if you want:
        # st.write(f"Processing line: {repr(line)}")
        if "\n" in line:
            line = line.replace("\n", "")
            doc.add_paragraph("")

        # 1) Page-break + vertical centered heading logic
        if "<break>" in line:
            doc.add_page_break()
            # look ahead for next non-empty, non-break
            j = i + 1
            while j < len(lines) and (lines[j].strip() == "" or "<break>" in lines[j]):
                j += 1
            if j < len(lines) and lines[j].strip().startswith("**") and lines[j].strip().endswith("**"):
                # add vertical spacing then center heading
                for _ in range(12):
                    doc.add_paragraph("")
                heading_text = lines[j].strip()[2:-2]
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(heading_text)
                run.bold = True
                set_paragraph_font(paragraph)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                i = j + 1
                in_break = False
                continue
            else:
                if j < len(lines):
                    for _ in range(12):
                        doc.add_paragraph("")
                    in_break = True
                else:
                    in_break = False
                i += 1
                continue

        # 2) Whole-line bold heading like **Heading**
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) >= 4:
            paragraph = doc.add_paragraph(stripped[2:-2])
            if paragraph.runs:
                paragraph.runs[0].bold = True
            if in_break:
                set_paragraph_font(paragraph)
            in_break = False
            i += 1
            continue

        # 3) '##' style heading
        if stripped.startswith("##"):
            paragraph = doc.add_paragraph(stripped.lstrip("#").strip())
            if paragraph.runs:
                paragraph.runs[0].bold = True
            if in_break:
                set_paragraph_font(paragraph)
            in_break = False
            i += 1
            continue

        # 4) Bullets: either indented (tabs) or normal '*'
        # Detect normalized tab-start or literal backslash-tab sequences
        normalized = line.replace("\\t", "\t")
        if normalized.lstrip().startswith("\t*") or normalized.lstrip().startswith("*") or stripped.startswith("*"):
            # Prefer to treat any line that contains leading tabs followed by * as bullet
            # Determine if there are tabs at start
            if normalized.startswith("\t") or normalized.lstrip().startswith("\t*") or ("\\t" in line and line.lstrip().startswith("\\t*")):
                # Indented bullet (could be one or many tabs)
                _handle_bullet_line(doc, line, bullet_symbols, in_break=in_break, cap_symbol=cap_bullet_symbol)
                in_break = False
                i += 1
                continue
            else:
                # Normal bullet starting with '*' (no leading tabs)
                # Use same cleaning logic as the helper: normalize then strip leading * chars
                without_tabs = line.replace("\\t", "\t").lstrip("\t ").lstrip()
                jdx = 0
                while jdx < len(without_tabs) and without_tabs[jdx] == "*":
                    jdx += 1
                clean_line = without_tabs[jdx:].lstrip()
                # create paragraph with first-level symbol
                p = doc.add_paragraph()
                if "**" in clean_line:
                    parts = re.split(r"(\*\*.*?\*\*)", clean_line)
                    p.add_run(f"{bullet_symbols[0]} ")
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            p.add_run(part[2:-2]).bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(f"{bullet_symbols[0]} {clean_line}")
                if in_break:
                    set_paragraph_font(p)
                in_break = False
                i += 1
                continue

        # 5) Inline bold fragments (non-heading, non-bullet)
        if "**" in line:
            p = doc.add_paragraph()
            bold_parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in bold_parts:
                if part.startswith("**") and part.endswith("**"):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
            if in_break:
                set_paragraph_font(p)
            in_break = False
            i += 1
            continue

        # 6) Indented text only (tabs but not bullets)
        if ("\t" in line and not line.lstrip().startswith("\t*")) or ("\\t" in line and not line.lstrip().startswith("\\t*")):
            tab_count = line.count("\t") if "\t" in line else line.count("\\t")
            clean_line = line.replace("\t", "").replace("\\t", "").strip()
            p = doc.add_paragraph(clean_line)
            p.paragraph_format.left_indent = Inches(0.3 * tab_count)
            if in_break:
                set_paragraph_font(p)
            in_break = False
            i += 1
            continue

        # 7) Default paragraph
        paragraph = doc.add_paragraph(stripped)
        if in_break:
            set_paragraph_font(paragraph)
        in_break = False
        i += 1
        continue

# download doc file from link
def generate_word_download_link(doc_data, filename):
    b64 = base64.b64encode(doc_data).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx">Click here to download the Word document</a>'
    return href

def generate_pdf_download_link(pdf_path, filename):
    with open(pdf_path, "rb") as f:
        pdf_data = f.read()
    b64 = base64.b64encode(pdf_data).decode()
    href = f'<a href="data:application/pdf;base64,{b64}" download="{filename}.pdf">Click here to download the PDF</a>'
    return href

# dynamic
def save_text_in_document_1(input,doc,flag,index=0,value=""):
    # print(text)
    
    style = doc.styles["Normal"]

    font = style.font
    font.name = "Times New Roman"
    
    if flag==0:
        process_text_to_docx(doc, input)

    elif flag==1:
        # num=value[0]
        image_extraction_paragraph = doc.add_paragraph(style='Normal')
        # run = image_extraction_paragraph.add_run(f"Image Extraction")
        # run.bold = True
        if input is not None:
        
            try:
                # st.image(input, caption="")
                doc.add_picture(input, width=Inches(5.7))
                
                image_extraction_paragraph = doc.add_paragraph(style='Normal')
                image_extraction_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = image_extraction_paragraph.add_run(f"{value}")
                run.bold = True
            except:
                image_text = "**Image is not found**"
                process_text_to_docx(doc, image_text)

                image_extraction_paragraph = doc.add_paragraph(style='Normal')
                image_extraction_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = image_extraction_paragraph.add_run(f"{value}")
                run.bold = True
                

        else:
            image_text = "Image is not found"
            process_text_to_docx(doc, image_text)


# -------------------------------Table add into output doc-------------------------------------------------------------------
    elif flag==2:
        i = 0
        
        for key,df in input.items():
            table_name=None
            # extract heading and table name from the table json
            try:
                table_name=df[0]["table_name"]
            except:
                pass

            table=df[0]["columns"]



         
            # Json table convert to df
            if len(table)>1:
                df0=pd.DataFrame(table)
            else:
                df0=pd.DataFrame(list(table[0].items()), columns=["Attribute", "Value"])

            

            # Add table name and heading into output documents
            try:
                if table_name is not None:

                        paragraph = doc.add_paragraph(f"{table_name.strip()}")
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run1 = paragraph.runs[0]  # Access the first run in the paragraph
                        run1.bold = True
                else:
                    paragraph = doc.add_paragraph(f"{table_name.strip()}")
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    run1 = paragraph.runs[0]  # Access the first run in the paragraph
                    run1.bold = True
            except:
                pass

            # Add table into output doc
            if df0.empty:
                print("The DataFrame is empty. No table will be created.")
            else:
                # Create table in document

                if "Value" not in df0.columns:

                    if "col_0" in df0.columns:
                        table = doc.add_table(rows=0, cols=len(df0.columns))
                
                    # Add rows from DataFrame
                        for index, row in df0.iterrows():
                            row_cells = table.add_row().cells
                            for j, cell in enumerate(row):
                                row_cells[j].text = str(cell)

                    else:
                        first_col_values = df0.iloc[:, 0].unique()
                        second_col_values = df0.iloc[:, 1].unique()
                        if len(first_col_values) == 1 and len(second_col_values) == 1 and len(df0.columns)>3:
                            first_2_column= df0.iloc[:, :2]
                            df0= df0.iloc[:, 2:]



                            table = doc.add_table(rows=0, cols=len(first_2_column.columns))

                            row_cells = table.add_row().cells
                            row_cells[0].text = str(first_2_column.columns[0])
                            row_cells[0].paragraphs[0].runs[0].font.bold = True
                            row_cells[1].text = str("\n".join(first_2_column.iloc[:, 0].unique()))

                            row_cells = table.add_row().cells
                            row_cells[0].text = str(first_2_column.columns[1])
                            row_cells[0].paragraphs[0].runs[0].font.bold = True
                            row_cells[1].text = str("\n".join(first_2_column.iloc[:, 1].unique()))

                            set_table_border(table)


                    
                        # Add rows from DataFrame
                        
                            

                        table = doc.add_table(rows=1, cols=len(df0.columns))

                        # Add column headers
                        hdr_cells = table.rows[0].cells
                        for j, col in enumerate(df0.columns):
                            if col is not None:  # Check for None values
                                hdr_cells[j].text = str(col)  # Ensure col is a string
                                hdr_cells[j].paragraphs[0].runs[0].font.bold = True
                                set_cell_background(hdr_cells[j], "F2F2F2")
                                

                        # Add rows from DataFrame
                        for index, row in df0.iterrows():
                            # Condition: More than 3 columns and more than 5 rows
                            if df0.shape[1] > 3 and df0.shape[0] > 5:
                                first_col_value = str(row[0]).strip().lower()
                                user_input_count = sum(1 for val in row if str(val).strip().lower() == "user input required")

                                # If first column is "user input required" OR count of "user input required" > 3 → skip
                                if first_col_value == "user input required" or user_input_count > 2:
                                    continue

                            # Add row to Word table
                            row_cells = table.add_row().cells
                            for j, cell in enumerate(row):
                                row_cells[j].text = str(cell)

         
                else:
                    table = doc.add_table(rows=0, cols=len(df0.columns))
                
                    # Add rows from DataFrame
                    for index, row in df0.iterrows():
                        row_cells = table.add_row().cells
                        for j, cell in enumerate(row):
                            row_cells[j].text = str(cell)

                set_table_border(table)
            doc.add_paragraph("")
            i += 1
    elif flag==3:
        # static
        pass
    else:
        # Web
        pass









