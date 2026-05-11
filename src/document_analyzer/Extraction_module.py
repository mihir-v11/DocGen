import streamlit as st
import json
import io
import os
import re
import shutil 
import logging
from docx import Document
from datetime import datetime
from docxcompose.composer import Composer
from dotenv import load_dotenv
from langchain.chat_models import AzureChatOpenAI
from openai import AzureOpenAI
from langchain.schema import SystemMessage, HumanMessage

from src.document_generate.dynamic_template import handle_user_message
from src.document_ingestion.data_collection import data_extraction
from src.document_analyzer.json_converter import extract_text_from_word
from src.document_generate.Assembling_appendix import assemble_appendices,convert_word_to_pdf
from src.document_analyzer.contents import refresh_toc_with_word,extract_headings_with_tables
from src.document_ingestion.paths import map_categories_to_json,map_categories_to_json_Executive_Summary,maping_folder

# Function to split the template text into two parts based on the section header
def Template_to_list(text):
    
    
    sections = [section.strip() for section in text.split('$') if section.strip()]
    split_index = None
    for i, item in enumerate(sections):
        if "DEVICE DESCRIPTION & PRODUCT SPECIFICATION".lower() in item.lower():
            split_index = i
            break

    if split_index is not None:
        part1 = sections[:split_index]
        part2 = sections[split_index:]
    else:
        part1 = sections
        part2 = []

    return part1, part2


# Convert the list of strings into a dictionary with unique keys
def convert_dict(list):

    result_dict = {}
    key_count = {}

    for item in list:
    # Split into key and value at first newline
        parts = item.split('\n', 1)
        if len(parts) == 2:
            key, value = parts[0].strip(), parts[1].strip()
        else:
            # If no newline found, use whole as key and empty value
            key, value = parts[0].strip(), ""

        # Handle duplicate keys by counting
        if key in key_count:
            key_count[key] += 1
            new_key = f"{key}_{key_count[key]}"
        else:
            key_count[key] = 1
            new_key = key

        result_dict[new_key] = value
    return result_dict



def extraction(template_file_path):

    load_dotenv()
    AZURE_KEY = os.getenv('AZURE_KEY')  # Replace with your actual key or keep it in .env
    AZURE_ENDPOINT = os.getenv('AZURE_ENDPOINT', 'https://api.geneai.thermofisher.com/dev/gpt4o')
    AZURE_NAME = os.getenv('AZURE_NAME', 'gpt-4o')
    AZURE_VERSION = os.getenv('AZURE_VERSION', '2024-05-01-preview')

    llm = None
    client = None
    try:
        llm = AzureChatOpenAI(
            deployment_name=AZURE_NAME,
            openai_api_key=AZURE_KEY,
            openai_api_base=AZURE_ENDPOINT,
            openai_api_version=AZURE_VERSION,
            openai_api_type="azure",
            temperature=0.1,
        )

        client = AzureOpenAI(
        api_key=AZURE_KEY,  
        api_version=AZURE_VERSION,
        azure_endpoint=AZURE_ENDPOINT,
)
    except Exception as e:
        print("Failed to initialize AzureChatOpenAI!")
        print("Error:", e)


                
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")


   
   
    # ------------------------------------- Devide json into 2 Part--------------------------------------------------------------------------------------------
    
    template_content = extract_text_from_word(template_file_path)
    print("1")
    # template convert into json
    part1,part2=Template_to_list(template_content)


    


    before_device_description = convert_dict(part1)
    after_device_description = convert_dict(part2)
    
    
    maping_folder()
    

                    
#--------------------------------------------------------------------------------------------------------------  
    # from pathlib import Path
    # current_dir = Path(__file__).resolve().parent
    # project_root = current_dir.parent.parent 
    
    # # 2. Define Folder Paths
    # template_dir = project_root / "templates"
    # output_dir = project_root / "data" / "artifacts" / "generated_output"
    
    # # 3. Create output directory if it doesn't exist
    # output_dir.mkdir(parents=True, exist_ok=True)
    
    # # 4. Define File Paths
    # template_path_1 = template_dir / "output_template_1.docx"
    # template_path = template_dir / "output_template.docx"
    # Output doc with footer header
    template_path_1 = fr"templates\\output_template_1.docx"
    template_path = fr"templates\\output_template.docx"
    
    # Create a new file name
    new_file_name = fr"DMF_Output_{timestamp}.docx"
    new_file_name_pdf = fr"DMF_Output_{timestamp}.pdf"
    # Load the template
    doc = Document(template_path_1)
    doc_w=Document(template_path_1)


    
    

# #     # --------------------------------Iterate over after executive Summary--------------------------------------------------------------------
    doc_1=doc
    # st.write("After Device Discription")
    for key,value in after_device_description.items():
            # Replace placeholders with actual values
            value = value.replace("[Device Name]", st.session_state.thermoDevice)
            # value = value.replace("[Short Name]", st.session_state.thermoDeviceShort)
            # value = value.replace("[Generic Name]",  st.session_state.thermoDeviceGeneric)
            
            flag=1
            index=1
            
            st.write(key)
            first_25 = value[:80]
            st.write(first_25)
            st.write("")
            st.write("")

            section_json={key:value}
            # if "image" in key.lower():
            # try:
            if "web" in key.lower():
                doc_1 = doc_w
                
            if "web" not in key.lower():
                # st.write("1")
                # path fetched for input files from path.py
                mapped_path = map_categories_to_json(section_json)
                # st.write("2")
                # st.write(mapped_path["data_path"])


                if mapped_path["data_path"]!="not found":
                    input_file_path=mapped_path["data_path"]
                    # st.write(input_file_path)
                    input_file_text=data_extraction(input_file_path)
                    
                    
                    response_data = handle_user_message(llm,client,key,value,doc_1,flag,index,input_file_text,input_file_path)
                else:
                    response_data = handle_user_message(llm,client,key,value,doc_1,flag,index)
            else:
                pass
                # response_data = handle_user_message(llm,client,key,value,doc_1,flag,index)  
            # except Exception as e:
            #     st.write("Error in processing the section:")
            #     st.write(f"An error occurred: {e}")
            #     continue
                
            # else:
            #     pass
            st.write("----------------------------------------------------------------------------------------------------------------------------")    


    # ------------------------------Iterate over Before executive Summary----------------------------------------------------------------
    
    doc1 = Document(template_path_1)

    st.write("----------------------------------------------------------------------------------------------------------------------------")
    
    for key,value in before_device_description.items():
            value = value.replace("[Device Name]", st.session_state.thermoDevice)
            # value = value.replace("[Short Name]", st.session_state.thermoDeviceShort)
            # value = value.replace("[Generic Name]",  st.session_state.thermoDeviceGeneric)
            
            flag=0
            index=0
            st.write(key)
            first_25 = value[:80]
            st.write(first_25)
            try:
                
                   
                    section_json={key:value}

                    # path fetched for input files from path.py
                    mapped_path = map_categories_to_json_Executive_Summary(section_json)
                    # st.write(mapped_path["data_path"])

                    if mapped_path["data_path"]!="not found" and len(mapped_path["data_path"]) > 0:
                        input_file_path=mapped_path["data_path"]
                        
                        
                       
                        
                        # extracted data from input file
                        input_file_text=data_extraction(input_file_path)
                    
                        response_data = handle_user_message(llm,client,key,value, doc1,flag,index,input_file_text,input_file_path)
                    else:
                                        
                        abc=""        

                        summary_input_text=extract_text_from_word(abc,doc)

                        response_data = handle_user_message(llm,client, key,value, doc1,flag,index,summary_input_text)
                    
            except:
                pass
            st.write("----------------------------------------------------------------------------------------------------------------------------")

    # Save the document to a new file----------------------------------------------------------------------------------------------------
    new_file_path = fr"data/artifacts/generated output file\\{new_file_name}" 
    new_file_path_temp =rf"data/artifacts/generated output file\Temp_{new_file_name}"
    final_document_doc = rf"data/artifacts/generated output file\\Final_{new_file_name}"
    new_file_path_pdf =rf"data/artifacts/generated output file\\{new_file_name_pdf}"
    final_document = fr"data/artifacts/generated output file\\Final_{new_file_name_pdf}.pdf" 



   # Add content and front page to the new document
    # doc0 = Document(template_path)
    
 
    # for paragraph in doc0.paragraphs:
    #     for run in paragraph.runs:
    #         if "[Device Name]" in run.text:
    #             # Replace the text
    #             run.text = run.text.replace("[Device Name]", st.session_state.thermoDevice)
    #             # Make the replaced text bold
    #             run.bold = True
    
    

    
    
   


    composer = Composer(doc1)
    composer.append(doc)
    composer.append(doc_w)
   
    composer.save(new_file_path_temp)


    doc_t=extract_headings_with_tables(new_file_path_temp, 0, final_document_doc)

    file_path_abs = os.path.abspath(final_document_doc)
    # st.write(file_path_abs)
    
    
    refresh_toc_with_word(file_path_abs)

    # doc123 = Document(final_document_doc)
    
    # doc_final = Document(template_path_1)

    # composer1 = Composer(doc_final)
    # composer1.append(doc0)
    # composer1.append(doc123)
    # composer1.save(new_file_path_temp)



   



    

    # for element in doc.element.body:
    #     doc1.element.body.append(element)
    
    # doc1.save(new_file_path)

    link_doc = Document(file_path_abs)
    output = io.BytesIO()
    link_doc.save(output)
    output.seek(0) 

#------------------------Appendices assembling---------------------------------------------------------
   
    # # Define the base folder path
    # base_folder = r"Extracted_folder"
    # folders = [f for f in os.listdir(base_folder) if os.path.isdir(os.path.join(base_folder, f))]
    # if len(folders)>1:
    #     base_folder = rf"Extracted_folder"
    # else:
    #     base_folder = rf"Extracted_folder/{folders[0]}"

    # # file_path_abs=r"C:\Users\mihir.vasoya\OneDrive - Thermo Fisher Scientific\Desktop\Git-Integration\RegulatoryDocGen\generated output file\DMF_86_11_5_2025.docx"
    # convert_word_to_pdf(file_path_abs,new_file_path_pdf)
    
    

    # res=assemble_appendices(base_folder, new_file_path_pdf, final_document)
    # print(res)


    return output,final_document,new_file_name_pdf






    # # final_document="Regulatory_Document_DMF_Output.pdf"
    # # new_file_name_pdf="Regulatory_Document_DMF_Output.pdf"


    # return final_document,new_file_name_pdf

    


