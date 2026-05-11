import re
import os
from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
import win32com.client as win32
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import pythoncom
import copy

def insert_toc(paragraph):
    """Insert a TOC field code at the given paragraph."""
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(fldSimple)


def has_page_break_before(para):
    """Check if paragraph has a page break before it (property or hard break)."""
    # Method 1: Check pageBreakBefore property
    if para._element.pPr is not None:
        pageBreakBefore = para._element.pPr.find(
            qn('w:pageBreakBefore')
        )
        if pageBreakBefore is not None:
            return True
    
    # Method 2: Check for hard page breaks in runs
    for run in para.runs:
        for br in run._element.findall(qn('w:br')):
            br_type = br.get(qn('w:type'))
            if br_type == 'page':
                return True
    
    return False


def has_page_break_in_previous_para(items, idx):
    """Check if previous paragraph ends with a hard page break."""
    if idx > 0:
        prev_item = items[idx - 1]
        if isinstance(prev_item, Paragraph):
            for run in prev_item.runs:
                for br in run._element.findall(qn('w:br')):
                    br_type = br.get(qn('w:type'))
                    if br_type == 'page':
                        return True
    return False


def copy_paragraph_with_format(source_para, target_doc, force_heading_level=None):
    """Copy paragraph text, formatting, and images from source to target with full format preservation."""
    
    # Preserve original style if not forcing heading
    if force_heading_level:
        target_para = target_doc.add_paragraph(style=f'Heading {force_heading_level}')
    else:
        # Try to preserve the original style
        try:
            target_para = target_doc.add_paragraph(style=source_para.style.name)
        except:
            target_para = target_doc.add_paragraph()
    
    # ALWAYS copy paragraph-level formatting (including alignment)
    if source_para.alignment is not None:
        target_para.alignment = source_para.alignment
    
    # Copy paragraph spacing
    if source_para.paragraph_format.space_before:
        target_para.paragraph_format.space_before = source_para.paragraph_format.space_before
    if source_para.paragraph_format.space_after:
        target_para.paragraph_format.space_after = source_para.paragraph_format.space_after
    if source_para.paragraph_format.line_spacing:
        target_para.paragraph_format.line_spacing = source_para.paragraph_format.line_spacing
    
    # Copy indentation
    if source_para.paragraph_format.left_indent:
        target_para.paragraph_format.left_indent = source_para.paragraph_format.left_indent
    if source_para.paragraph_format.right_indent:
        target_para.paragraph_format.right_indent = source_para.paragraph_format.right_indent
    if source_para.paragraph_format.first_line_indent:
        target_para.paragraph_format.first_line_indent = source_para.paragraph_format.first_line_indent
 
    # Copy text and run-level formatting
    for run in source_para.runs:
        new_run = target_para.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        
        # Copy font properties
        if run.font.size:
            new_run.font.size = run.font.size
        if run.font.name:
            new_run.font.name = run.font.name
        
        # Copy font color
        if run.font.color.rgb:
            new_run.font.color.rgb = run.font.color.rgb
 
    # Copy inline images
    for inline_shape in source_para._element.xpath('.//w:drawing'):
        for pic in inline_shape.xpath('.//a:blip'):
            rId = pic.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if rId:
                try:
                    image_part = source_para.part.related_parts[rId]
                    image_stream = BytesIO(image_part.blob)
                    target_para.add_run().add_picture(image_stream, width=Inches(4.0))
                except:
                    pass
    
    return target_para

def copy_table_with_format(source_table, target_doc,source_doc,index):
    """Copy a table from source to target document while preserving formatting, merged cells, and structure."""
    doc2 = Document(source_doc)

    # Get the target table from the template
    target_table = doc2.tables[index]
    # Insert the table into the current document
    body = target_doc.element.body
    paragraphs = target_doc.paragraphs
    insert_index = len(paragraphs)-1  # Insert after the last paragraph by default

    if insert_index < len(paragraphs):
        para = paragraphs[insert_index]._element
        body.insert(body.index(para)+1, target_table._element)
   
def iter_paragraphs_and_tables(doc):
    """Yield each paragraph and table from the document body in order."""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)

def extract_headings_with_tables(input_path, flag=1, output_path=None):
    """Extract content and add TOC while preserving formatting."""
    if flag == 1:
        source_doc = input_path 
    else:
        source_doc = Document(input_path)

    new_doc = Document()

    # TOC Title
    toc_title = new_doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run("Contents")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # TOC Field
    insert_toc(new_doc.add_paragraph())
    new_doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    
    items = list(iter_paragraphs_and_tables(source_doc))
    
    heading_pattern = re.compile(r'^(\d+(\.0|\.))([\s\.:)])')
    sub_heading_pattern = re.compile(r'^\d+\.\d+')
    index=0
   
    for idx, block in enumerate(items):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            
            # Check BOTH paragraph property AND hard page breaks
            has_break_before = has_page_break_before(block)
            # has_break_in_prev = has_page_break_in_previous_para(items, idx)
            
            # Add page break if detected
            if has_break_before: # or has_break_in_prev:
                new_doc.add_page_break()
            
            # Now copy the paragraph content
            match = heading_pattern.match(text)
            if match:
                copy_paragraph_with_format(block, new_doc, force_heading_level=1)
            elif sub_heading_pattern.match(text):
                copy_paragraph_with_format(block, new_doc, force_heading_level=2)
            else:
                copy_paragraph_with_format(block, new_doc)
        
        elif isinstance(block, Table):
            copy_table_with_format(block, new_doc, input_path, index)
            index += 1

    if output_path is not None:
        new_doc.save(output_path)
    
    return new_doc

def refresh_toc_with_word(doc_path):
    """Refresh TOC using Microsoft Word COM."""
    pythoncom.CoInitialize()
    
    try:
        word = win32.gencache.EnsureDispatch("Word.Application")
        word.Visible = False
        
        abs_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_path)

        # Update all TOCs
        for toc in doc.TablesOfContents:
            toc.Update()
            toc.Range.Font.Bold = True

        doc.Save()
        doc.Close()
        print("✅ TOC updated and saved.")
        
    except Exception as e:
        print(f"❌ Error updating TOC: {e}")
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

# Example usage
# if __name__ == "__main__":
#     input_file = r"data\artifacts\generated output file\CER_Text_part_3_11_2025.docx"
#     output_file = r"final_document_doc10.docx"
    
#     # Extract and add TOC
#     doc = extract_headings_with_tables(input_file, flag=0, output_path=output_file)
    
#     # Refresh TOC
#     refresh_toc_with_word(output_file)