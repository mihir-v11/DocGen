import streamlit as st
import os
import io
import logging
import time
import shutil

# from Features.Summarization import summarized_Document
# from Features.scrap import scrapping
from src.document_analyzer.Extraction_module import extraction
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import zipfile
import stat
import gc
import pandas as pd
from src.document_generate.doc_generate import generate_word_download_link

# streamlit run app.py --server.maxUploadSize=1000

# Reset logging configuration
for handler in logging.root.handlers[:]:
    logging.root.removeHandler(handler)

logger = logging.getLogger()

# Check if handlers are already present
if not logger.hasHandlers():
    # Set logging level
    logger.setLevel(logging.INFO)

    # Create handlers
    file_handler = logging.FileHandler('logs/app.log')
    console_handler = logging.StreamHandler()

    # Set levels for handlers
    file_handler.setLevel(logging.INFO)
    console_handler.setLevel(logging.INFO)

    # Create formatter and add to handlers
    formatter = logging.Formatter('%(asctime)s - %(levelname)s - %(message)s')
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)

    # Add handlers to logger
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
# logging.basicConfig(filename='app.log', level=logging.INFO, filemode='a',
#                     format='%(asctime)s - %(levelname)s - %(message)s')
logging.info("App started")
def add_logo_to_docx(doc_data, logo_path="pdf_logo.png", text="The world leader in serving science"):
    """Adds a logo to the left side of the top header section and specified text to the right side.

    Args:
        doc_data (bytes): The content of the Word document as a byte stream.
        logo_path (str, optional): The path to the logo image file. Defaults to "pdf_logo.png".
        text (str, optional): The text to add to the header. Defaults to "world leader in serving science".

    Returns:
        bytes: The modified Word document content as a byte stream.
    """
    
    document = Document(doc_data)

    # Get the first section's header
    section = document.sections[0]
    header = section.header

    # Create a table with two cells: one for the logo, one for the text
    table = header.add_table(rows=1, cols=2, width=Inches(6))
    table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Align table to left

    # Adjust column widths
    table.columns[0].width = Inches(1.5)  # Width for the logo
    table.columns[1].width = Inches(4.5)  # Width for the text

    # Add logo to the first cell
    logo_cell = table.cell(0, 0)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(logo_path, width=Inches(1.25))

    # Add text to the second cell, aligning it to the right
    text_cell = table.cell(0, 1)
    text_paragraph = text_cell.paragraphs[0]
    text_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Create a run with a manual line break
    bold_run = text_paragraph.add_run("The world leader\nin serving science")
    bold_run.font.name = "Times New Roman"
    bold_run.bold = True  # Make the text bold

    # Optionally set margins for all sections
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Save the document and return as byte stream
    document.save("temp.docx")
    with open("temp.docx", "rb") as f:
        return f.read()


def cleanup_environment():
    """
    Clean up all temporary variables, cached data, and free up memory for the next run.
    """
    print("Cleaning up the environment...")

    # Clear global variables
    global all_tokens, competitor_analysis
    all_tokens = []
    competitor_analysis = []

    # Clear any other global variables or temporary data
    globals_to_clear = [var for var in globals() if not var.startswith("__") and not callable(globals()[var])]
    for var in globals_to_clear:
        if isinstance(globals()[var], (pd.DataFrame, list, dict, set)):
            globals()[var] = None  # Reset to None
            print(f"Cleared variable: {var}")

    # Force garbage collection
    gc.collect()
    print("Environment cleanup completed.")


# Title of the app
st.markdown(
    "<h1 style='text-align: center; color: red;'>Document Generator</h1>",
    unsafe_allow_html=True
)
st.markdown("""
    <style>
    /* Styling for the main submit button */
    div.stButton > button {
        background-color: red;  /* Red background */
        width: 200px;           /* Adjust width */
        height: 40px;           /* Adjust height */
        border: none;           /* Remove border */
        cursor: pointer;        /* Change cursor on hover */
        border-radius: 25px;    /* Rounded corners */
        color: Black;           /* White text color */
        border: 2px solid white; /* Border */
        margin-top: 5px;
    }

    /* Hover effect for the buttons */
    div.stButton > button:hover {
        background-color: white;
        color: red;
        border: 2px solid red;
    }
    </style>
""", unsafe_allow_html=True)

image_name=[]
# Example: Add your logo (local or online image path)
# logo_url = r"static\logo1.png"  # Replace with your logo URL or local path

# Add the logo to the sidebar
# st.sidebar.image(logo_url, width=200,caption=None )

# Initialize session state variables if they don't exist
if 'selected_feature' not in st.session_state:
    st.session_state.selected_feature = None
    
col1, col2, col3 = st.columns(3)


#---------------------------------------Side Bar------------------------------------------------------------------------------

st.sidebar.header("Select Option")
with st.sidebar:
    
    if st.button("Document Generator"):
        st.session_state.selected_feature = 'Technical'

    # if st.button("CER File"):
    #     st.session_state.selected_feature = 'scraping'

    # if st.button("Plant Master File"):
    #     st.session_state.selected_feature = 'summarization'

#---------------------------------------Side Bar Ends------------------------------------------------------------------------------




def clear_extracted_folder(folder_path):
    """
    Clears all files and subdirectories in the specified folder.
    """
    if os.path.exists(folder_path):
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.unlink(file_path)  # Remove file or symbolic link
                elif os.path.isdir(file_path):
                    # Use shutil.rmtree with error handling for permissions
                    shutil.rmtree(file_path, onerror=handle_remove_readonly)
            except Exception as e:
                print(f"Failed to delete {file_path}. Reason: {e}")
    else:
        os.makedirs(folder_path)
 

def handle_remove_readonly(func, path, exc_info):
    """
    Handle read-only files or folders during deletion.
    """
    try:
        # Change the file or folder's permissions to writable
        os.chmod(path, stat.S_IWUSR)
        func(path)
    except Exception as e:
        print(f"Failed to modify permissions for {path}. Reason: {e}")


# Technical File
if st.session_state.selected_feature == 'Technical':
    st.subheader("📄 File Generator")
    st.write("")

    uploaded_file = st.file_uploader("Upload a ZIP file containing your folder", type="zip")
   

    # st.session_state.technical_data_sheet = st.file_uploader("Upload Technical Data Sheet", type="pdf", key="tech_data")
    
    # Text inputs
    st.session_state.thermoDevice = st.text_input("Enter Device Name", value="")
    # st.session_state.thermoDeviceGeneric = st.text_input("Enter Thermo Device Generic Name", value="")
    # st.session_state.thermoDeviceShort = st.text_input("Enter Thermo Device Family Short Name", value="")

    # prev_gen_available = st.radio(
    #     "Previous Generation Device is:",
    #     ("Available", "Not Available"),
    #     key="prev_gen_availability"
    # )
    
    # Initialize variables
    # st.session_state.previousGenDevice = None
    # st.session_state.prev_gen_data_sheet = None
    
    # Show/hide inputs based on radio button selection
    # if prev_gen_available == "Available":
    #     st.session_state.previousGenDevice = st.text_input("Enter Previous Generation Device Name", value="")
    #     st.session_state.prev_gen_data_sheet = st.file_uploader("Upload Previous Generation Data Sheet", type="pdf", key="prev_gen_data")

    
    # st.session_state.user_keyword = st.text_input("Enter User Keyword", value="")
    # st.session_state.refined_keyword = st.text_input("Enter Refined Keyword", value="")

    

    # template_doc= r"templates\DMR_Template_21-5-2025.docx"
    # template_doc=r"templates\template_1 - Copy.docx"
    from pathlib import Path

    BASE_DIR = Path(__file__).resolve().parents[2]
    template_doc = BASE_DIR / "template" / "DMR_Template_18_12_2025.docx"
    # template_doc=r"templates\DMR_Template_18_12_2025.docx"
        
    token_info=[] 
    input_tokens = 0
    output_tokens = 0
    total_tokens = 0
  
    st.write("")
    st.write("")

    if st.button("Submit"):
            clear_extracted_folder(r"data\artifacts\Extracted_folder")
            if (
                uploaded_file is None
                # st.session_state.technical_data_sheet is None or
                # st.session_state.prev_gen_data_sheet is None or
                # not st.session_state.thermoDevice.strip() or
                # not st.session_state.previousGenDevice.strip() or
                # not st.session_state.user_keyword.strip() and
                # not st.session_state.refined_keyword.strip()
            ):
                st.warning("🚨 Please fill in all required fields and upload files before submitting.")
            
            
            
            else:
                # Extract the uploaded Zip files in Extracted_folder
                if uploaded_file:
                    with open("temp.zip", "wb") as f:
                        f.write(uploaded_file.getbuffer())
                
                    # Extract the ZIP file
                    extract_dir = r"data\\artifacts\\Extracted_folder"
                    # Clean up previous extractions
                    with zipfile.ZipFile("temp.zip", "r") as zip_ref:
                        zip_ref.extractall(extract_dir)

                    folder_structure = {}
                
                    # Recursively walk through the extracted folder
                    
                    for root, dirs, files in os.walk(extract_dir):
                        # Store the relative folder path
                        relative_root = os.path.relpath(root, extract_dir)
                        if relative_root == ".":
                            relative_root = "Root"  # Rename the top-level folder for clarity
                        
                        # Store the folder and its files in the dictionary
                        folder_structure[relative_root] = files
                        
                        # Display files in this folder, if any
                        if files:
                            for file in files:
                                relative_path = os.path.join(relative_root, file)
                        else:
                            pass
                    
                    # Store the folder structure in session state for later use
                    st.session_state['folder_structure'] = folder_structure
                    st.session_state['extract_dir'] = extract_dir
                    os.remove("temp.zip")
                #--------------------------------------------------------------------------------------




                start_time = time.time()

                # pdf_path,pdf_name = extraction(template_doc)

                output,pdf_path,pdf_name = extraction(template_doc)

                # pdf_path="ajsagdjgas"
                response_time = time.time() - start_time
                # response_time=1455



                logging.info(f"Total Tokens for Doc Generate------------------------------------------")
                logging.info(f"Input Tokens: {input_tokens}")
                logging.info(f"Output Tokens: {output_tokens}")
                logging.info(f"Total Tokens: {total_tokens}")
                logging.info(f"Response generation time: {response_time:.2f} seconds")

                # file_name = os.path.splitext(documnet_name)[0]
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            
                prefix = "Regulatory_Document" 
                # final_file_name = f"{prefix}_{documnet_name}_{timestamp}"

                if pdf_path: 
                    st.subheader("📥 Please download the generated document.")

                    st.html(f"<p><strong>Response generation time:- </strong>{response_time:.2f} seconds</p>")
                
                # st.html(f"<p><strong>PDF Path:- </strong>Desktop\Git-Integration\RegulatoryDocGen\{pdf_path}</p>")

                #     # shutil.rmtree("Extracted_folder")
                #     if 'folder_structure' in st.session_state:
                #         del st.session_state['folder_structure']
                #     if 'extract_dir' in st.session_state:
                #         del st.session_state['extract_dir']
                clear_extracted_folder(r"data\artifacts\Extracted_folder")
                cleanup_environment()

                final_file_name = f"DMF_Final_{timestamp}.docx"
                st.markdown(f"DMF Document 👉 {generate_word_download_link(output.getvalue(),final_file_name)}", unsafe_allow_html=True)   
